#!/usr/bin/env python3
"""
Build specialized DOCX packets for printing.

Outputs:
- Flashcards DOCX: 8 cards/page, duplex-friendly front/back pages.
- Practice exams DOCX: question block, choices block, answer key on its own page.

These DOCX files intentionally differ from the vault Markdown packets. Markdown remains
informational/reference-oriented; DOCX is print-layout-oriented.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parents[1]
ARTIFACT_DIR = REPO_ROOT / "08 Printable Study Materials" / "Build Artifacts"
FLASHCARD_OUT = ARTIFACT_DIR / "Water-Operator-Vault-Flashcards-Packet.docx"
EXAM_OUT = ARTIFACT_DIR / "Water-Operator-Vault-Practice-Exam-Packet.docx"

FLASHCARD_FILES = [
    "02 Flash Cards/T5 Flash Cards - PFAS and MCLs.md",
    "02 Flash Cards/T5 Flash Cards - SDWA and Regulations.md",
    "02 Flash Cards/T5 Flash Cards - Water Quality and Contaminants.md",
    "02 Flash Cards/T5 Flash Cards - Water Math.md",
    "02 Flash Cards/T5 Flash Cards - Disinfection and CT.md",
    "02 Flash Cards/T5 Flash Cards - Treatment Processes.md",
    "02 Flash Cards/T5 Flash Cards - Lab Procedures and Methods.md",
    "02 Flash Cards/T5 Flash Cards - Cross Connection Control.md",
    "02 Flash Cards/T5 Flash Cards - Distribution Source Hydrants and Disinfection.md",
]

EXAM_FILES = [
    "03 Practice Exams/Randomized Final Drafts/Practice-Exam-01-Core-Regulations-MCLs-and-Operator-Math - Randomized Final Draft.md",
    "03 Practice Exams/Randomized Final Drafts/Practice-Exam-02-Treatment-Processes-and-Water-Quality - Randomized Final Draft.md",
    "03 Practice Exams/Randomized Final Drafts/Practice-Exam-03-Source-Water-Distribution-Sampling-and-Public-Health - Randomized Final Draft.md",
    "03 Practice Exams/Randomized Final Drafts/Practice-Exam-04-Advanced-Math-Pumps-SCADA-and-Operations - Randomized Final Draft.md",
    "03 Practice Exams/Randomized Final Drafts/Practice-Exam-05-Scenario-Judgment-Compliance-and-Troubleshooting - Randomized Final Draft.md",
    "03 Practice Exams/Randomized Final Drafts/T5-Math-Only-Practice-Bank - Randomized Final Draft.md",
    "03 Practice Exams/Randomized Final Drafts/T5-PFAS-and-Emerging-Contaminants-Mini-Exam - Randomized Final Draft.md",
    "03 Practice Exams/Randomized Final Drafts/T5-Practice-Exam-06-Advanced-Scenario-Mix - Randomized Final Draft.md",
    "03 Practice Exams/Randomized Final Drafts/T5-Lab-Procedures-and-Methods-Mini-Exam - Randomized Final Draft.md",
    "03 Practice Exams/Randomized Final Drafts/T5-Cross-Connection-Control-Mini-Exam - Randomized Final Draft.md",
    "03 Practice Exams/Randomized Final Drafts/T5-Distribution-Source-Hydrants-and-Disinfection-Mini-Exam - Randomized Final Draft.md",
]

FRONT_RE = re.compile(r"^\*\*Front:\*\*\s*(?P<front>.*)$")
BACK_RE = re.compile(r"^\*\*Back:\*\*\s*(?P<back>.*)$")
QUESTION_RE = re.compile(
    r"^(?P<num>\d+)\.\s+(?P<prompt>.*?)\s+"
    r"A\)\s+(?P<A>.*?)\s+"
    r"B\)\s+(?P<B>.*?)\s+"
    r"C\)\s+(?P<C>.*?)\s+"
    r"D\)\s+(?P<D>.*)$"
)
KEY_RE = re.compile(r"^\|\s*(?P<num>\d+)\s*\|\s*(?P<answer>[ABCD])\s*\|\s*(?P<review>.*?)\s*\|")
BATCH_RE = re.compile(r"Batch ID:\s*`(?P<batch>[^`]+)`|batch_id:\s*(?P<batch2>\S+)")


@dataclass
class FlashCard:
    deck: str
    number: int
    front: str
    back: str
    source: str


@dataclass
class ExamQuestion:
    num: int
    prompt: str
    choices: dict[str, str]
    answer: str = ""
    review: str = ""


@dataclass
class Exam:
    title: str
    batch_id: str
    questions: list[ExamQuestion]


def strip_markdown(text: str) -> str:
    text = re.sub(r"\[\[([^#\]|]+)#?([^\]|]*)\|?([^\]]*)\]\]", lambda m: m.group(3) or m.group(1), text)
    text = re.sub(r"\[\[([^\]]+)\]\]", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.strip()


def deck_title(text: str, fallback: str) -> str:
    for line in text.splitlines():
        if line.startswith("# "):
            return strip_markdown(line[2:])
    return fallback


def parse_source_from_back(back: str, default_source: str) -> tuple[str, str]:
    if "Source:" in back:
        answer, source = back.split("Source:", 1)
        return strip_markdown(answer), strip_markdown(source)
    return strip_markdown(back), default_source


def parse_flashcards() -> list[FlashCard]:
    cards: list[FlashCard] = []
    for rel in FLASHCARD_FILES:
        path = REPO_ROOT / rel
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8")
        title = deck_title(text, path.stem)
        default_source = f"Deck: {title}"
        current_num = 0
        pending_front: str | None = None
        for line in text.splitlines():
            if line.startswith("### Card"):
                match = re.search(r"Card\s+(\d+)", line)
                current_num = int(match.group(1)) if match else current_num + 1
                pending_front = None
                continue
            front = FRONT_RE.match(line)
            if front:
                pending_front = strip_markdown(front.group("front"))
                continue
            back = BACK_RE.match(line)
            if back and pending_front:
                answer, source = parse_source_from_back(back.group("back"), default_source)
                cards.append(FlashCard(title, current_num, pending_front, answer, source))
                pending_front = None
    return cards


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="777777", size="8", visible: bool = True) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single" if visible else "nil")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, lines: list[tuple[str, bool]], font_size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for idx, (text, bold) in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)


def set_document_margins(doc: Document, top=0.25, bottom=0.25, left=0.25, right=0.25) -> None:
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
        section.header_distance = Inches(0.0)
        section.footer_distance = Inches(0.0)


def set_landscape(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)


def add_card_grid(doc: Document, cards: list[FlashCard], side: str) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        row.height = Inches(1.75)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = Inches(5.2)
            set_cell_border(cell, color="999999", size="6")
            set_cell_shading(cell, "FFFFFF")
    for idx, card in enumerate(cards):
        row = idx // 2
        col = 1 - (idx % 2) if side == "back" else idx % 2
        cell = table.cell(row, col)
        if side == "front":
            lines = [
                (f"FRONT | {card.deck} | Card {card.number}", True),
                ("", False),
                (card.front, False),
            ]
        else:
            lines = [
                (f"BACK | Card {card.number}", True),
                (card.back, False),
                ("", False),
                (f"Source: {card.source}", False),
            ]
        set_cell_text(cell, lines, font_size=8)


def build_flashcards_docx() -> None:
    cards = parse_flashcards()
    doc = Document()
    set_landscape(doc)
    set_document_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(8)

    for start in range(0, len(cards), 8):
        chunk = cards[start : start + 8]
        while len(chunk) < 8:
            chunk.append(FlashCard("", 0, "", "", ""))
        add_card_grid(doc, chunk, "front")
        doc.add_page_break()
        add_card_grid(doc, chunk, "back")
        if start + 8 < len(cards):
            doc.add_page_break()

    ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(FLASHCARD_OUT)
    print(f"Wrote {FLASHCARD_OUT} with {len(cards)} cards")


def parse_answer_key(lines: list[str]) -> dict[int, tuple[str, str]]:
    key: dict[int, tuple[str, str]] = {}
    in_key = False
    for line in lines:
        if line.startswith("## Answer Key Page"):
            in_key = True
            continue
        if in_key and line.startswith("## Metadata"):
            break
        if in_key:
            match = KEY_RE.match(line)
            if match:
                key[int(match.group("num"))] = (match.group("answer"), strip_markdown(match.group("review")))
    return key


def parse_exam(path: Path) -> Exam:
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = path.stem.replace(" - Randomized Final Draft", "")
    for line in lines:
        if line.startswith("# "):
            title = strip_markdown(line[2:].replace(" - Randomized Final Draft", ""))
            break
    batch = ""
    for match in BATCH_RE.finditer(text):
        batch = match.group("batch") or match.group("batch2") or ""
        if batch:
            break
    key = parse_answer_key(lines)
    questions: list[ExamQuestion] = []
    in_questions = False
    for line in lines:
        if line.startswith("## Questions"):
            in_questions = True
            continue
        if in_questions and line.startswith("---"):
            break
        if not in_questions:
            continue
        match = QUESTION_RE.match(line)
        if not match:
            continue
        num = int(match.group("num"))
        answer, review = key.get(num, ("", ""))
        questions.append(
            ExamQuestion(
                num=num,
                prompt=strip_markdown(match.group("prompt")),
                choices={letter: strip_markdown(match.group(letter)) for letter in "ABCD"},
                answer=answer,
                review=review,
            )
        )
    return Exam(title=title, batch_id=batch, questions=questions)


def add_exam_question(doc: Document, q: ExamQuestion) -> None:
    q_para = doc.add_paragraph()
    q_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    q_para.paragraph_format.space_before = Pt(6)
    q_para.paragraph_format.space_after = Pt(3)
    q_para.paragraph_format.line_spacing = 1.05
    label = q_para.add_run(f"Question {q.num}. ")
    label.bold = True
    label.font.size = Pt(10)
    text = q_para.add_run(q.prompt)
    text.font.size = Pt(10)

    choice_table = doc.add_table(rows=2, cols=2)
    choice_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    choice_table.autofit = True
    choice_order = [("A", 0, 0), ("B", 0, 1), ("C", 1, 0), ("D", 1, 1)]
    for letter, row, col in choice_order:
        cell = choice_table.cell(row, col)
        set_cell_border(cell, visible=False)
        set_cell_text(cell, [(f"{letter}) {q.choices[letter]}", False)], font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()


def build_exam_docx() -> None:
    exams = [parse_exam(REPO_ROOT / rel) for rel in EXAM_FILES if (REPO_ROOT / rel).exists()]
    doc = Document()
    set_document_margins(doc, top=0.5, bottom=0.5, left=0.55, right=0.55)
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Water Operator Vault Practice Exam Packet")
    run.bold = True
    run.font.size = Pt(16)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("Questions are left-aligned. Choices are arranged in two columns and two rows. Answer keys begin on separate pages.").font.size = Pt(9)
    doc.add_page_break()

    for exam_index, exam in enumerate(exams):
        heading = doc.add_paragraph()
        heading.style = doc.styles["Heading 1"]
        heading.add_run(exam.title)
        batch = doc.add_paragraph()
        batch.add_run(f"Batch ID: {exam.batch_id}").bold = True
        for q in exam.questions:
            add_exam_question(doc, q)
        doc.add_page_break()
        key_heading = doc.add_paragraph()
        key_heading.style = doc.styles["Heading 1"]
        key_heading.add_run(f"Answer Key - {exam.title}")
        key_batch = doc.add_paragraph()
        key_batch.add_run(f"Batch ID: {exam.batch_id}").bold = True
        key_table = doc.add_table(rows=1, cols=3)
        key_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Q", "Answer", "Review"]
        for i, header in enumerate(headers):
            cell = key_table.cell(0, i)
            set_cell_shading(cell, "D9EAF7")
            set_cell_border(cell)
            set_cell_text(cell, [(header, True)], font_size=9)
        for q in exam.questions:
            row = key_table.add_row()
            values = [str(q.num), q.answer, q.review or "source-supported with caveats"]
            for i, value in enumerate(values):
                cell = row.cells[i]
                set_cell_border(cell, color="999999", size="4")
                set_cell_text(cell, [(value, False)], font_size=8)
        if exam_index + 1 < len(exams):
            doc.add_page_break()

    ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(EXAM_OUT)
    print(f"Wrote {EXAM_OUT} with {sum(len(e.questions) for e in exams)} questions")


def main() -> None:
    build_flashcards_docx()
    build_exam_docx()


if __name__ == "__main__":
    main()
